from PIL import Image
import docx
from docx.shared import Inches,Cm
import os


def generate_docx(filepaths, preset):
    doc = docx.Document()

    for sec in doc.sections:
        sec.top_margin = Cm(0.5)
        sec.bottom_margin = Cm(0.5)
        sec.left_margin = Cm(0.5)
        sec.right_margin = Cm(0.5)

    para = doc.add_paragraph()
    run = para.add_run()


    exists=False
    for file in filepaths:
        img = Image.open(file)
        k = img.width/img.height
        if k<1:
            temp = img.transpose(Image.ROTATE_90)
            temp.save("temp.jpg")
            file = "temp.jpg"
            k = temp.width/temp.height
            exists = True
        try:
            run.add_picture(file,width=Inches(preset*k), height=Inches(preset))
        except docx.image.exceptions.UnrecognizedImageError:
            img.save("unrecognised_image.jpg")
            file = "unrecognised_image.jpg"
            run.add_picture(file,width=Inches(preset*k), height=Inches(preset))
            os.remove("unrecognised_image.jpg")

        run.add_text(" ")
        if exists:
            os.remove("temp.jpg")
            exists = False

    doc.save("PicAlign_output.docx")



